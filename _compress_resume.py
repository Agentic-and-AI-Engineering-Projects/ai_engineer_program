"""
Compress a pandoc-generated .docx resume to fit 2 / 2.5 pages.
Does NOT change any text — only:
  - Margins (1.0 in → 0.55 in)
  - Body font (11pt → 10pt)
  - Heading sizes (slightly smaller)
  - Paragraph spacing (less before/after)
  - Line spacing (slightly tighter)
"""
import sys
from docx import Document
from docx.shared import Inches, Pt


def compress(path: str) -> None:
    doc = Document(path)

    # 1. Tighten margins.
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # 2. Modify the base "Normal" style — body text.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.10

    # 3. Heading sizes (smaller) + tighter spacing.
    heading_specs = {
        "Heading 1": (16, 6, 2),
        "Heading 2": (13, 5, 2),
        "Heading 3": (11, 4, 1),
        "Heading 4": (10, 3, 1),
        "Title":     (18, 0, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        try:
            h = doc.styles[name]
            h.font.size = Pt(size)
            h.paragraph_format.space_before = Pt(before)
            h.paragraph_format.space_after = Pt(after)
        except KeyError:
            pass

    # 4. List item + pandoc-specific styles.
    for name in [
        "List Bullet", "List Paragraph",
        "Compact",               # pandoc tight-list paragraph style
        "First Paragraph",       # pandoc first-para after heading
        "Body Text",
    ]:
        try:
            s = doc.styles[name]
            s.font.size = Pt(10)
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(1)
            s.paragraph_format.line_spacing = 1.10
        except KeyError:
            pass

    # 5. Walk every paragraph and enforce caps on spacing (without
    #    overriding heading styles, which are larger by design).
    for para in doc.paragraphs:
        pf = para.paragraph_format
        if pf.space_before is not None and pf.space_before.pt > 6:
            pf.space_before = Pt(2)
        if pf.space_after is not None and pf.space_after.pt > 4:
            pf.space_after = Pt(2)
        if pf.line_spacing is not None and pf.line_spacing > 1.2:
            pf.line_spacing = 1.10

    # 6. Tables (if any) — tighter cell padding.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(1)
                    pf.line_spacing = 1.10

    doc.save(path)


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "RESUME_DC_INFRASTRUCTURE_PM.docx"
    compress(target)
    print(f"Compressed: {target}")
